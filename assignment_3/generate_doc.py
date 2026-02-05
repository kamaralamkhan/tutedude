from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    doc = Document()

    # Title
    title = doc.add_heading('Flask & MongoDB Integration - Assignment 3', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Info
    p = doc.add_paragraph()
    p.add_run('Assignment: ').bold = True
    p.add_run('Flask API & MongoDB Atlas\n')
    p.add_run('Date: ').bold = True
    p.add_run('February 5, 2026\n')
    p.add_run('GitHub Repo: ').bold = True
    p.add_run('https://github.com/kamaralamkhan/tutedude')

    def add_task(title, command, explanation, output):
        doc.add_heading(title, level=1)
        doc.add_paragraph(explanation)
        
        # Code/Command box
        p_cmd = doc.add_paragraph()
        run_cmd = p_cmd.add_run(command)
        run_cmd.font.name = 'Courier New'
        run_cmd.font.size = Pt(10)
        p_cmd.paragraph_format.left_indent = Pt(20)

        # Output box
        doc.add_paragraph('Program Output/Result:', style='Normal').runs[0].bold = True
        p_out = doc.add_paragraph()
        run_out = p_out.add_run(output)
        run_out.font.name = 'Courier New'
        run_out.font.size = Pt(9)
        run_out.font.color.rgb = RGBColor(0, 51, 102) # Dark blue
        p_out.paragraph_format.left_indent = Pt(20)

    # Task 1
    add_task(
        '1. Flask API Route (/api)',
        '@app.route("/api")\ndef get_api_data():\n    with open("data.json", "r") as file:\n        data = json.load(file)\n    return jsonify(data)',
        'Created a Flask route that reads experimental data from a local data.json file and returns it as a JSON response.',
        '[\n  {"id": 1, "name": "Python Basics", "status": "Completed"},\n  {"id": 2, "name": "Flask Web Dev", "status": "In Progress"}\n]'
    )

    # Task 2
    add_task(
        '2. MongoDB Atlas Integration',
        'client = MongoClient(MONGO_URI)\ncollection = db.submissions\n\n# Inserting Data\ncollection.insert_one({"name": name, "email": email, "message": message})',
        'Implemented a registration form that captures User details and inserts them directly into a MongoDB Atlas cloud database. The app handles success redirects and database errors.',
        'Data submitted successfully\n(Redirected to /success)'
    )

    doc.save('Assignment_3_Submission.docx')
    print("Assignment 3 Document created successfully.")

if __name__ == "__main__":
    create_doc()
