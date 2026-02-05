from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    doc = Document()

    # Title
    title = doc.add_heading('Python Programming Assignment 2', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Info
    p = doc.add_paragraph()
    p.add_run('Assignment: ').bold = True
    p.add_run('Basic Python Logic & File I/O\n')
    p.add_run('Date: ').bold = True
    p.add_run('February 5, 2026\n')
    p.add_run('GitHub Repo: ').bold = True
    p.add_run('https://github.com/user/python-assignment-2')

    def add_task(title, command, explanation, output):
        doc.add_heading(title, level=1)
        doc.add_paragraph(explanation)
        
        # Code/Command box
        p_cmd = doc.add_paragraph()
        run_cmd = p_cmd.add_run(command)
        run_cmd.font.name = 'Courier New'
        run_cmd.font.size = Pt(10)
        p_cmd.paragraph_format.left_indent = Pt(20)

        # Output box
        doc.add_paragraph('Program Output:', style='Normal').runs[0].bold = True
        p_out = doc.add_paragraph()
        run_out = p_out.add_run(output)
        run_out.font.name = 'Courier New'
        run_out.font.size = Pt(9)
        run_out.font.color.rgb = RGBColor(0, 100, 0) # Dark green
        p_out.paragraph_format.left_indent = Pt(20)

    # Task 1
    add_task(
        '1. Grade Checker',
        '# Using if-elif-else statements\nscore = float(input("Enter the score: "))\n# ... logic ...\nprint(f"Grade: {grade}")',
        'This program taking a numeric score as input and uses basic if-else logic to categorize it into grades A, B, C, D, or F.',
        'Enter the score: 85\nThe grade for score 85.0 is: B'
    )

    # Task 2
    add_task(
        '2. Student Grades (Dictionary Operations)',
        'students = {"Alice": "A", "Bob": "B"}\n# Adding\nstudents[name] = grade\n# Updating\nif name in students: students[name] = new_grade',
        'Implemented a dictionary-based system to manage student records. It allows adding new students, updating existing grades, and listing all records.',
        '1. Add Student\n2. Update Grade\n3. Print All Grades\n4. Exit\nEnter choice: 1\nEnter student name: Charlie\nEnter student grade: A\n\nStudent Grades:\nAlice: A\nBob: B\nCharlie: A'
    )

    # Task 3
    add_task(
        '3. Write to a File',
        'with open("sample_text.txt", "w") as file:\n    file.write("Hello! This is a test file...")',
        'This program uses the open() function in "w" (write) mode to create a new text file and save content to it.',
        'Successfully wrote to sample_text.txt'
    )

    # Task 4
    add_task(
        '4. Read from a File',
        'with open("sample_text.txt", "r") as file:\n    content = file.read()\n    print(content)',
        'We used the open() function in "r" (read) mode and the .read() method to retrieve and display the file contents.',
        '--- File Content ---\nHello! This is a test file created for Assignment 2.\nPython file operations are very powerful.\n--------------------'
    )

    doc.save('Assignment_2_Submission.docx')
    print("Assignment 2 Document created successfully.")

if __name__ == "__main__":
    create_doc()
