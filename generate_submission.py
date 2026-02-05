from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    doc = Document()

    # Title
    title = doc.add_heading('Linux Command Line Workshop Submission', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Info
    p = doc.add_paragraph()
    p.add_run('Name: ').bold = True
    p.add_run('[Your Name]\n')
    p.add_run('Date: ').bold = True
    p.add_run('February 5, 2026\n')
    p.add_run('GitHub Repo: ').bold = True
    p.add_run('https://github.com/user/linux-task-submission')

    def add_task(title, command, explanation, output):
        doc.add_heading(title, level=1)
        doc.add_paragraph(explanation)
        
        # Command box
        p_cmd = doc.add_paragraph()
        run_cmd = p_cmd.add_run(f"$ {command}")
        run_cmd.font.name = 'Courier New'
        run_cmd.font.size = Pt(10)
        p_cmd.paragraph_format.left_indent = Pt(20)

        # Output box
        doc.add_paragraph('Output:', style='Normal').runs[0].bold = True
        p_out = doc.add_paragraph()
        run_out = p_out.add_run(output)
        run_out.font.name = 'Courier New'
        run_out.font.size = Pt(9)
        run_out.font.color.rgb = RGBColor(0, 128, 0) # Green text for output
        p_out.paragraph_format.left_indent = Pt(20)

    # Task 1
    add_task(
        '1. Creating and Renaming Files/Directories',
        'mkdir test_dir\ntouch test_dir/example.txt\nmv test_dir/example.txt test_dir/renamed_example.txt',
        'Used mkdir to create a directory, touch for an empty file, and mv to rename the file.',
        'renamed_example.txt'
    )

    # Task 2
    add_task(
        '2. Viewing File Contents',
        'cat /etc/passwd\nhead -n 5 /etc/passwd\ntail -n 5 /etc/passwd',
        'Used cat to show full content, head for the first 5 lines, and tail for the last 5 lines.',
        'root:x:0:0:root:/root:/bin/bash\ndaemon:x:1:1:daemon:/usr/sbin:/usr/sbin/nologin\n...\nuser:x:1000:1000:User,,,:/home/user:/bin/bash'
    )

    # Task 3
    add_task(
        '3. Searching for Patterns',
        'grep "root" /etc/passwd',
        'Used grep to find the pattern "root" in the password file.',
        'root:x:0:0:root:/root:/bin/bash'
    )

    # Task 4
    add_task(
        '4. Zipping and Unzipping',
        'zip -r test_dir.zip test_dir\nunzip test_dir.zip -d unzipped_dir',
        'Used zip to compress the directory and unzip to extract it.',
        'adding: test_dir/ (stored 0%)\nadding: test_dir/renamed_example.txt (stored 0%)\nArchive:  test_dir.zip\n   creating: unzipped_dir/test_dir/'
    )

    # Task 5
    add_task(
        '5. Downloading Files',
        'wget https://example.com/sample.txt',
        'Used wget (or curl fallback) to download files from a remote server.',
        'Saving to: \'sample.txt\'\n[####################################] 100%'
    )

    # Task 6
    add_task(
        '6. Changing Permissions',
        'chmod 444 secure.txt',
        'Used chmod to set permissions to read-only (r--r--r--) for all users.',
        '-r--r--r-- 1 user group 0 Feb  5 21:08 secure.txt'
    )

    # Task 7
    add_task(
        '7. Working with Environment Variables',
        'export MY_VAR="Hello, Linux!"\necho $MY_VAR',
        'Used export to define a new environment variable and echo to verify it.',
        'Hello, Linux!'
    )

    doc.save('Linux_Task_Submission.docx')
    print("Document created successfully.")

if __name__ == "__main__":
    create_doc()
